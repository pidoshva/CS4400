from docx import Document

# Load the cloc report from the text file
with open("cloc_report.txt", "r") as file:
    cloc_data = file.readlines()

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading("Lines of Code Report", level=1)

# Add the cloc data to the document
doc.add_paragraph("Generated using cloc tool\n")
doc.add_paragraph("Report Details:\n")

# Write the lines from the cloc report
for line in cloc_data:
    doc.add_paragraph(line.strip())

# Save the document
doc.save("cloc_report.docx")
